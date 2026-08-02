#!/usr/bin/env python3
"""Convert knowledge-base Markdown documents to .docx, keeping tables intact.

Why tables matter more than looks here: these files are price lists, and the
price extractor reads *tables* (price_tables._from_docx). A converter that
flattened a Markdown table into paragraphs would keep the document readable
while silently making every price row unextractable, so pipe tables become
real Word tables and nothing else in the file is allowed to disturb them.

Usage:
    python scripts/md_to_docx.py SRC_DIR [DEST_DIR]

SRC_DIR is walked recursively; each *.md becomes *.docx beside it (or under
DEST_DIR, mirroring the relative path). Existing .docx files are overwritten.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_FENCE_RE = re.compile(r"^\s*```")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
_HR_RE = re.compile(r"^\s*([-*_])\s*(\1\s*){2,}$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|[\s:|-]*$")

# Inline runs: **bold**, *italic*, `code`, [text](url) -> plain text.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]*\))")

# The source Markdown was produced by a PDF→MD converter that emitted LaTeX
# for symbols and units ("Thép cuộn $\\phi$ 6", "$\\mathrm{đồng/m^3}$").
# Cleaning it at conversion time means the .docx a reader opens is already
# readable, instead of only being cleaned later during price extraction.
_LATEX_CMD_RE = re.compile(r"\\(?:mathrm|mathbf|mathit|text|textbf|rm)\s*\{([^{}]*)\}")
_LATEX_SYMBOLS = {
    r"\phi": "Ø",
    r"\Phi": "Ø",
    r"\times": "x",
    r"\pm": "±",
    r"\le": "≤",
    r"\ge": "≥",
    r"\%": "%",
}
_SUPERSCRIPT_RE = re.compile(r"\^\{?(\d)\}?")


def _strip_latex(text: str) -> str:
    if "$" not in text and "\\" not in text:
        return text
    out = _LATEX_CMD_RE.sub(r"\1", text)
    for cmd, repl in _LATEX_SYMBOLS.items():
        out = out.replace(cmd, repl)
    out = _SUPERSCRIPT_RE.sub(r"\1", out)
    return re.sub(r"\s+", " ", out.replace("$", "")).strip()


_LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]*)\)$")


def _add_runs(paragraph, text: str) -> None:
    """Render inline markdown into runs. Unknown syntax falls through as
    literal text rather than being dropped — losing content silently would be
    worse than a stray asterisk."""
    text = _strip_latex(text)
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif (m := _LINK_RE.match(part)) is not None:
            label, url = m.group(1), m.group(2)
            paragraph.add_run(label).underline = True
            if url and url not in label:
                paragraph.add_run(f" ({url})").font.size = Pt(8)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    # "<br>" separates stacked values inside one cell; a newline keeps them
    # visible without splitting the cell (which would shift every column).
    return [_strip_latex(re.sub(r"<br\s*/?>", "\n", c).strip()) for c in s.split("|")]


def _add_table(doc: Document, block: list[str]) -> None:
    rows = [_split_row(x) for x in block if not _TABLE_SEP_RE.match(x)]
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_runs(para, row[j] if j < len(row) else "")
            if i == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()

    table_block: list[str] = []
    code_block: list[str] = []
    in_code = False

    def flush_table() -> None:
        nonlocal table_block
        if len(table_block) >= 2:
            _add_table(doc, table_block)
        elif table_block:
            for raw in table_block:
                _add_runs(doc.add_paragraph(), raw)
        table_block = []

    def flush_code() -> None:
        nonlocal code_block
        if code_block:
            para = doc.add_paragraph()
            run = para.add_run("\n".join(code_block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            code_block = []

    for line in lines:
        if _FENCE_RE.match(line):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_block.append(line)
            continue

        if line.lstrip().startswith("|"):
            table_block.append(line)
            continue
        flush_table()

        if not line.strip():
            continue
        if _HR_RE.match(line):
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if (m := _HEADING_RE.match(line)) is not None:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 6))
            continue
        if (m := _QUOTE_RE.match(line)) is not None:
            para = doc.add_paragraph(style="Intense Quote")
            _add_runs(para, m.group(1))
            continue
        if (m := _BULLET_RE.match(line)) is not None:
            _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            continue
        if (m := _ORDERED_RE.match(line)) is not None:
            _add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            continue
        _add_runs(doc.add_paragraph(), line.strip())

    if in_code:
        flush_code()
    flush_table()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    if len(sys.argv) < 2:
        print(__doc__)
        return 2
    src = Path(sys.argv[1])
    dest = Path(sys.argv[2]) if len(sys.argv) > 2 else src
    md_files = sorted(src.rglob("*.md"))
    if not md_files:
        print(f"Không tìm thấy .md nào trong {src}")
        return 1
    for md in md_files:
        out = (dest / md.relative_to(src)).with_suffix(".docx")
        convert(md, out)
        print(f"{md.name}  ->  {out.name}")
    print(f"\nĐã chuyển {len(md_files)} file.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
