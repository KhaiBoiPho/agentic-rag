"""DOCX chunker — preserves document order: text paragraphs + embedded tables.

Strategy (RAGFlow-inspired):
  Iterate the DOCX body in DOM order, collecting (text | table) elements.
  Tables are rendered as HTML. Context is attached from surrounding text.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from app.core.chunking.base import BaseChunker, add_table_context, count_tokens, naive_merge
from app.core.chunking.models import Chunk, ChunkType


class DocxChunker(BaseChunker):
    def chunk(self, filename: str, content: bytes, **kwargs) -> list[Chunk]:
        doc_id = kwargs.get("document_id", "")
        kb_id = kwargs.get("kb_id", "")

        doc = Document(BytesIO(content))
        ordered = self._extract_ordered(doc)
        merged = self._merge_and_contextualize(ordered)

        chunks: list[Chunk] = []
        for item in merged:
            c = Chunk(
                content=item["text"],
                chunk_type=item["chunk_type"],
                document_id=doc_id,
                kb_id=kb_id,
                filename=filename,
                context_above=item.get("context_above", ""),
                context_below=item.get("context_below", ""),
                token_count=count_tokens(item["text"]),
            )
            chunks.append(c)

        return chunks

    def _extract_ordered(self, doc: Document) -> list[dict]:
        """Walk body XML, yielding paragraphs and tables in document order."""
        items: list[dict] = []
        body = doc.element.body

        for child in body.iterchildren():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    items.append({"text": text, "chunk_type": ChunkType.TEXT})

            elif tag == "tbl":
                table = DocxTable(child, doc)
                html = self._table_to_html(table)
                items.append({"text": html, "chunk_type": ChunkType.TABLE})

        return items

    def _table_to_html(self, table: DocxTable) -> str:
        html = ["<table>"]
        for i, row in enumerate(table.rows):
            cells = row.cells
            tag = "th" if i == 0 else "td"
            row_html = "".join(f"<{tag}>{c.text.strip()}</{tag}>" for c in cells)
            html.append(f"<tr>{row_html}</tr>")
        html.append("</table>")
        return "\n".join(html)

    def _merge_and_contextualize(self, items: list[dict]) -> list[dict]:
        result: list[dict] = []
        text_buffer: list[str] = []

        def flush():
            if not text_buffer:
                return
            for t in naive_merge(
                text_buffer,
                chunk_token_num=self.chunk_token_num,
                delimiter=self.delimiter,
                overlap_percent=self.overlap_percent,
            ):
                result.append({"text": t, "chunk_type": ChunkType.TEXT})
            text_buffer.clear()

        for item in items:
            if item["chunk_type"] == ChunkType.TEXT:
                text_buffer.append(item["text"])
            else:
                flush()
                result.append(item)

        flush()
        add_table_context(result, self.table_context_size)
        return result
